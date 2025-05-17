from os import path
from json import dump
from docx import Document

from solution.docx_converter.models import Message, Chat, CompanyChat
from cleaner import clear_text, REMOVE_CHARS


# Сперва необходимо загрузить .docx файл в переменную объекта .docx
def load_docx_from_project(filename: str) -> tuple[str, Document]:
    if not filename.endswith('.docx'):
        raise Exception(f"Некорректный формат {filename}.\nОжидался .docx")

    try:
        document = Document(filename)
    except Exception as e:
        print(f"Не удалось загрузить {filename}: {e}")
        return filename, None

    # path.basename возвращает последнее поле пути + убирает формат файла
    name_only = path.basename(filename)
    return name_only, document


# Далее необходимо извлечь данные из .docx объекта и поместить их в переменную класса Chat
# Для простоты последующей обработки
def convert_from_docx(data: tuple[str, Document]) -> Chat:
    if not data[1].tables:
        raise ValueError(f"Файл {data[0]} не содержит таблиц")

    name = data[0][:-5]
    table = data[1].tables[0]

    numbers: list[str] = []
    messages: list[Message] = []

    for row in table.rows[1:]:
        cells = row.cells
        message = Message(sender=cells[1].text, text=cells[2].text)

        numbers.append(cells[0].text)
        messages.append(message)

    chat = Chat(
        name=name,
        numbers=set(numbers),
        messages=messages
    )

    return chat


# Далее логично обработать переменную класса Chat в переменную упрощенного текстового формата
# CompanyChat, содержащую информацию в формате просто диалога от отправителя к отправителю
def format_chat(chat: Chat) -> CompanyChat:
    dialog = ""

    for message in chat.messages:
        cleared_text = clear_text(message.text, REMOVE_CHARS)
        if cleared_text.strip():
            dialog += (message.sender + ': ' + '\n' +
                       cleared_text + '\n\n')

    name_data = CompanyChat(company=chat.name, whole_chat=dialog)
    return name_data


# Наконец нужно записать готовые данные из переменной CompanyChat в .json файл. Запись будет
# производиться отдельно - для каждой компании свой файл.
# (Учел то, что данные обычно берутся не в формате .docx)
def write_chat_to_json(
    chat: CompanyChat,
    folder_path: str
) -> None:

    if not chat.company or not chat.whole_chat:
        raise ValueError("Пустые имя компании или текст")

    file_path = path.join(folder_path, chat.company + ".json")
    chat_json = {"company": chat.company, "chat": chat.whole_chat}

    with open(file_path, "w", encoding="utf-8") as f:
        dump(chat_json, f, ensure_ascii=False, indent=4)

    return None
